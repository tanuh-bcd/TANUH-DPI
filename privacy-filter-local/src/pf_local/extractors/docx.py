"""DOCX extractor + redactor.

Reads paragraph + table cell text, runs detection on the joined text, then
rewrites the document in-place by replacing matched spans with
[REDACTED:LABEL] tokens. Preserves paragraph structure (formatting on
specific runs may be lost where a span crosses runs — acceptable for a
testing app).
"""
from __future__ import annotations

from pathlib import Path
from typing import List, Dict, Any

from docx import Document


_SEP = "\n"


def _iter_blocks(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def extract_text(path: Path) -> str:
    doc = Document(str(path))
    return _SEP.join(p.text for p in _iter_blocks(doc))


def redact(path: Path, entities: List[Dict[str, Any]], out_path: Path) -> None:
    doc = Document(str(path))
    paragraphs = list(_iter_blocks(doc))

    # Build a flat mapping: char offset -> paragraph index.
    text_pieces = [p.text for p in paragraphs]
    offsets: List[int] = []
    cursor = 0
    for piece in text_pieces:
        offsets.append(cursor)
        cursor += len(piece) + len(_SEP)

    spans = sorted(
        [e for e in entities if e.get("start") is not None and e.get("end") is not None],
        key=lambda e: e["start"],
    )

    # Apply per-paragraph: for each paragraph, collect spans inside its range.
    for i, p in enumerate(paragraphs):
        p_start = offsets[i]
        p_end = p_start + len(p.text)
        local_spans = [
            (e["start"] - p_start, e["end"] - p_start, e.get("entity_group", "PII"))
            for e in spans
            if e["start"] >= p_start and e["end"] <= p_end
        ]
        if not local_spans:
            continue

        new_text = p.text
        # Apply right-to-left so offsets remain valid.
        for s, t, label in sorted(local_spans, key=lambda x: x[0], reverse=True):
            new_text = new_text[:s] + f"[REDACTED:{label.upper()}]" + new_text[t:]

        # Replace the runs: clear existing runs, write new text in run[0].
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = new_text
        else:
            p.add_run(new_text)

    doc.save(str(out_path))
